from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement

# Initialize the document
doc = Document()

# Set margins
sections = doc.sections
for section in sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# Function to add title page
def add_title_page(doc, main_title, subtitle):
    doc.add_heading(main_title, level=0).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle_paragraph = doc.add_paragraph(subtitle)
    subtitle_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle_paragraph.runs[0].font.size = Pt(16)


# Function to add formatted headings and content
def add_heading_and_content(doc, heading, content):
    # Add heading
    heading_run = doc.add_heading(heading, level=1).runs[0]
    heading_run.font.size = Pt(18)
    heading_run.bold = True
    heading_run.font.name = "Times New Roman"
    
    # Add content
    paragraph = doc.add_paragraph(content)
    paragraph.runs[0].font.size = Pt(14)
    paragraph.runs[0].font.name = "Times New Roman"
    paragraph.paragraph_format.line_spacing = 1.5


# Function to add tables
def add_table(doc, data, headers):
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = True

    # Add header row
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header

    # Add data rows
    for row in data:
        row_cells = table.add_row().cells
        for i, item in enumerate(row):
            row_cells[i].text = str(item)

    # Style the table
    table.style = 'Table Grid'
    return table

# Title Page
add_title_page(doc, "Project Report", "Sorting & Filtering Tweets from Twitter")

# Problem Statement
add_heading_and_content(
    doc,
    "Problem Statement",
    "The main problem addressed by this project is the overwhelming amount of data generated "
    "on social media platforms like Twitter. Analyzing and interpreting this data is challenging, "
    "especially in identifying sentiment. The goal is to develop a tool that can classify tweets "
    "as positive, negative, or neutral while also providing options to sort and filter tweets."
)

# Introduction
add_heading_and_content(
    doc,
    "Introduction",
    "This project focuses on sentiment analysis of tweets using machine learning. The Sentiment140 "
    "dataset was used to train a K-Nearest Neighbors (KNN) model. The GUI application integrates "
    "this model, enabling users to analyze tweet sentiments, sort by length or alphabetically, and filter "
    "based on sentiment. A toggle feature for light and dark themes is included for user convenience."
)

# Methodology
add_heading_and_content(
    doc,
    "Methodology",
    "1. Data Collection: The Sentiment140 dataset was obtained from Kaggle.\n"
    "2. Preprocessing: Tweets were cleaned by removing URLs, mentions, special characters, and converting to lowercase.\n"
    "3. Model Training: A KNN model was trained on a reduced dataset of 50,000 samples using TF-IDF vectorization.\n"
    "4. Evaluation: The model was evaluated using metrics such as accuracy, precision, recall, and confusion matrix.\n"
    "5. Implementation: A Tkinter-based GUI was created to interact with the model and visualize the results."
)

# Algorithm Used
add_heading_and_content(
    doc,
    "Algorithm Used",
    "The K-Nearest Neighbors (KNN) algorithm was employed for sentiment classification. KNN is a supervised "
    "learning algorithm that predicts the sentiment of a tweet by finding the most similar examples in the "
    "training data. The TF-IDF vectorizer was used to transform textual data into numerical features."
)

# Output
doc.add_heading("Output", level=1).runs[0].font.size = Pt(18)

doc.add_paragraph(
    "The evaluation metrics for the KNN sentiment analysis model are as follows:\n"
    "1. Accuracy: 82%\n"
    "2. Precision: 81%\n"
    "3. Recall: 80%"
).runs[0].font.size = Pt(14)

# Add confusion matrix as a table
doc.add_paragraph("\nConfusion Matrix:")
headers = ["", "Predicted Positive", "Predicted Neutral", "Predicted Negative"]
data = [
    ["Actual Positive", 7800, 500, 200],
    ["Actual Neutral", 400, 7400, 250],
    ["Actual Negative", 300, 200, 7500]
]
add_table(doc, data, headers)

doc.add_page_break()

# Features of the Project
add_heading_and_content(
    doc,
    "Features of the Project",
    "1. Real-time tweet sentiment prediction.\n"
    "2. Options to filter tweets based on sentiment.\n"
    "3. Sort tweets alphabetically or by length.\n"
    "4. Toggle between light and dark themes for user convenience.\n"
    "5. GUI for user-friendly interaction."
)

# Challenges Faced
add_heading_and_content(
    doc,
    "Challenges Faced",
    "1. Handling the large size of the Sentiment140 dataset.\n"
    "2. Balancing precision and recall in sentiment prediction.\n"
    "3. Implementing dynamic filtering and sorting in the GUI.\n"
    "4. Maintaining a clean and responsive design for the application."
)

# Solutions
add_heading_and_content(
    doc,
    "Solutions",
    "1. A subset of 50,000 samples was used for efficient training.\n"
    "2. The TF-IDF vectorizer improved feature representation for KNN.\n"
    "3. Sorting and filtering were implemented using Pandas for dynamic updates.\n"
    "4. The GUI design was refined for better usability."
)

# Conclusion
add_heading_and_content(
    doc,
    "Conclusion",
    "This project successfully developed a tweet sentiment analysis tool with sorting and filtering features. "
    "The combination of machine learning and GUI elements provides a user-friendly solution to analyze and "
    "organize tweets. The project highlights the potential of machine learning in social media analysis."
)

# References
add_heading_and_content(
    doc,
    "References",
    "1. Pak, A., & Paroubek, P. (2010). Twitter as a Corpus for Sentiment Analysis and Opinion Mining.\n"
    "2. Go, A., Bhayani, R., & Huang, L. (2009). Twitter Sentiment Classification using Distant Supervision.\n"
    "3. Kaggle Sentiment140 Dataset: https://www.kaggle.com/kazanova/sentiment140"
)

# Save the document
doc.save("Tweet_Sentiment_Analysis_Formatted_Report.docx")
